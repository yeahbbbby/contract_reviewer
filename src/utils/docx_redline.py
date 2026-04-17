"""
Word红线修订工具
生成带修订痕迹的Word文档
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict
from datetime import datetime


class DocxRedline:
    """Word红线修订工具"""

    def __init__(self, original_path: str):
        """
        初始化红线工具

        Args:
            original_path: 原始文档路径
        """
        self.doc = Document(original_path)
        self.author = "合同审查系统"
        self.date = datetime.now()

    def add_deletion(self, paragraph_index: int, text: str):
        """
        添加删除标记

        Args:
            paragraph_index: 段落索引
            text: 要删除的文本
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return

        para = self.doc.paragraphs[paragraph_index]

        # 查找并标记删除
        if text in para.text:
            # 创建删除标记
            run = para.add_run()
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:author'), self.author)
            del_elem.set(qn('w:date'), self.date.isoformat())

            del_run = OxmlElement('w:r')
            del_text = OxmlElement('w:delText')
            del_text.text = text

            del_run.append(del_text)
            del_elem.append(del_run)
            run._element.append(del_elem)

    def add_insertion(self, paragraph_index: int, text: str, position: int = -1):
        """
        添加插入标记

        Args:
            paragraph_index: 段落索引
            text: 要插入的文本
            position: 插入位置（-1表示末尾）
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return

        para = self.doc.paragraphs[paragraph_index]

        # 创建插入标记
        run = para.add_run()
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:author'), self.author)
        ins_elem.set(qn('w:date'), self.date.isoformat())

        ins_run = OxmlElement('w:r')
        ins_text = OxmlElement('w:t')
        ins_text.text = text

        ins_run.append(ins_text)
        ins_elem.append(ins_run)
        run._element.append(ins_elem)

    def add_comment(self, paragraph_index: int, comment_text: str):
        """
        添加批注

        Args:
            paragraph_index: 段落索引
            comment_text: 批注内容
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return

        para = self.doc.paragraphs[paragraph_index]

        # 简化版：在段落末尾添加批注标记
        run = para.add_run(f" [批注: {comment_text}]")
        run.font.color.rgb = None  # 红色
        run.font.size = None

    def apply_modifications(self, modifications: List[Dict]):
        """
        批量应用修改

        Args:
            modifications: 修改列表
        """
        for mod in modifications:
            mod_type = mod.get('type')
            para_index = mod.get('paragraph_index', 0)

            if mod_type == 'delete':
                self.add_deletion(para_index, mod.get('text', ''))
            elif mod_type == 'insert':
                self.add_insertion(para_index, mod.get('text', ''))
            elif mod_type == 'comment':
                self.add_comment(para_index, mod.get('text', ''))
            elif mod_type == 'replace':
                self.add_deletion(para_index, mod.get('old_text', ''))
                self.add_insertion(para_index, mod.get('new_text', ''))

    def add_review_summary(self, summary: Dict):
        """
        在文档开头添加审查摘要

        Args:
            summary: 审查摘要
        """
        # 在文档开头插入新段落
        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run("=" * 50).bold = True

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run("合同审查摘要").bold = True
        para.alignment = 1  # 居中

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run("=" * 50).bold = True

        # 添加摘要内容
        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run(f"审查日期: {summary.get('date', '')}")

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run(f"合规评分: {summary.get('compliance_score', 0)}/100")

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run(f"可行性结论: {summary.get('feasibility', '')}")

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run(f"识别风险: {summary.get('risk_count', 0)}项")

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run(f"修改建议: {summary.get('modification_count', 0)}项")

        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.add_run("")

    def highlight_risky_clauses(self, risky_clauses: List[Dict]):
        """
        高亮风险条款

        Args:
            risky_clauses: 风险条款列表
        """
        for clause in risky_clauses:
            keyword = clause.get('keyword', '')
            risk_level = clause.get('risk_level', 'medium')

            # 查找包含关键词的段落
            for i, para in enumerate(self.doc.paragraphs):
                if keyword in para.text:
                    # 添加批注
                    comment = f"[{risk_level.upper()}风险] {clause.get('description', '')}"
                    self.add_comment(i, comment)

    def save(self, output_path: str):
        """
        保存文档

        Args:
            output_path: 输出路径
        """
        self.doc.save(output_path)

    def generate_redline_from_report(self, final_report: Dict, output_path: str):
        """
        根据审查报告生成红线版本

        Args:
            final_report: 最终审查报告
            output_path: 输出路径
        """
        # 添加审查摘要
        summary = final_report.get('structured_report', {}).get('summary', {})
        summary['date'] = datetime.now().strftime('%Y年%m月%d日')
        summary['modification_count'] = len(final_report.get('modifications', []))
        self.add_review_summary(summary)

        # 应用修改建议（转换为具体的修改操作）
        modifications = self._convert_modifications_to_operations(
            final_report.get('modifications', [])
        )
        self.apply_modifications(modifications)

        # 保存
        self.save(output_path)

    def _convert_modifications_to_operations(self, modifications: List[Dict]) -> List[Dict]:
        """
        将修改建议转换为具体的修改操作

        Args:
            modifications: 修改建议列表

        Returns:
            修改操作列表
        """
        operations = []

        for mod in modifications:
            description = mod.get('description', '')
            priority = mod.get('priority', 'medium')

            # 简化处理：将所有建议作为批注添加
            # 实际应用中需要更复杂的逻辑来定位具体段落
            operations.append({
                'type': 'comment',
                'paragraph_index': 0,  # 需要智能定位
                'text': f"[{priority.upper()}] {description}"
            })

        return operations
